"""
utils.py — File extraction helpers for PDF and DOCX resumes.
"""
import io


def extract_text_from_pdf(uploaded_file) -> str:
    try:
        file_bytes = uploaded_file.read()
    except Exception:
        file_bytes = uploaded_file.getvalue()

    if not file_bytes:
        return ""

    # Method 1: pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages = [p.extract_text() or "" for p in pdf.pages]
        text = "\n".join(pages).strip()
        if text:
            return text
    except Exception:
        pass

    # Method 2: PyPDF2
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n".join(pages).strip()
        if text:
            return text
    except Exception:
        pass

    # Method 3: PyMuPDF
    try:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = [page.get_text() for page in doc]
        text = "\n".join(pages).strip()
        if text:
            return text
    except Exception:
        pass

    return ""


def extract_text_from_docx(uploaded_file) -> str:
    try:
        file_bytes = uploaded_file.read()
    except Exception:
        file_bytes = uploaded_file.getvalue()

    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts).strip()
    except Exception:
        return ""
